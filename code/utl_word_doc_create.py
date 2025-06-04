

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import matplotlib.pyplot as plt
import os

class CreateWordDoc():
	def __init__(self, doc_path, font_size = 11, font_name = 'Times New Roman'):
		self.doc_path = doc_path
		self.doc = Document()

		# Access the 'Normal' style
		#normal_style = self.doc.styles['Normal']
		# Create a custom style by copying the 'Normal' style
		#self.custom_style = self.doc.styles.add_style('CustomStyle', 1)  # 1 indicates a paragraph style
		#self.custom_style.base_style = normal_style  # Link the new style to 'Normal'

		# Explicitly set font properties for the custom style
		#self.custom_font = self.custom_style.font
		#self.custom_font.name = font_name             # Change font name
		#self.custom_font.size = font_size             # Change font size


	def add_text_line_new_paragraph(self, txt):
	# Add a text line
		self.doc.add_paragraph(txt)
	
	def add_text_line_last_paragraph(self, txt):
		last_paragraph = self.doc.paragraphs[-1]
	# Add a new line within the same paragraph
		last_paragraph.add_run("\n" + txt)
	
	def add_text_with_highlight(self, txt, position_lst):
		position_lst2 = [num - 1 for num in position_lst]  # Adjust for 1-based index
		paragraph = self.doc.add_paragraph()
		# Add styled text
		for i, char in enumerate(txt):
			run = paragraph.add_run(char)
			if i in position_lst2:  
				run.bold = True
				run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
				run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Yellow background
			run.font.size = Pt(11)
			#run.font.name = self.custom_font.name

	def add_picture(self,img_path, imag=''):
		if imag == '':
			self.doc.add_picture(img_path, width=Pt(400))

	def save_doc(self):
		self.doc.save(self.doc_path)

	
	


if __name__ == '__main__':
	plt.plot([1, 2, 3], [4, 5, 6])
	plt.title("Sample Plot")
	image_path = "plot.png"
	plt.savefig(image_path)
	plt.close()
	doc_path = 'C:/_Shaul/Python/_Shani/output/output.docx'
	doc1 = CreateWordDoc(doc_path, 16, 'Times New Roman')
	doc1.add_text_line_new_paragraph(f'Hellow word - printing a simple text')
	string1 = '1234567890_' * 30
	doc1.add_text_with_highlight(string1,[1,3,5,7,200,202, 330, 400])
	doc1.add_picture(image_path)
	os.remove(image_path)  # Clean up the generated image file
	for i in range(10):
		doc1.add_text_line_new_paragraph(f'{i} - Adding line new paragraph')
	for i in range(10):
		doc1.add_text_line_last_paragraph(f'{i} - Adding line last paragraph')
	doc1.save_doc()

	print(f"Document saved as {doc_path}")
	
